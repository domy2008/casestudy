"""Unit tests for the per-format document loaders (task 6.2).

Covers the nominal extraction path for every supported format (DOCX, XLSX,
TXT, Markdown, and PDF) with fixture files created at runtime, asserting that
body text and table row/column structure survive extraction (Req 5.2), plus
the error paths: a corrupt file raises a parse failure (Req 5.4) and an
unsupported extension routes to the unsupported-format error.

reportlab is not a pinned dependency, so PDF fixtures are generated with a
tiny hand-built PDF that pypdf/pdfplumber can read. Hand-built PDFs carry no
ruled table geometry, so the PDF *table* assertion is skipped with a clear
reason while PDF text extraction and loader dispatch are still exercised.
"""

from __future__ import annotations

import importlib.util
from pathlib import Path

import pytest

from app.core.models import ExtractedContent
from app.kb.loaders import (
    DocumentParseError,
    DocxLoader,
    Loader,
    MarkdownLoader,
    PdfLoader,
    SUPPORTED_EXTENSIONS,
    TextLoader,
    UnsupportedFormatError,
    XlsxLoader,
    find_loader,
    get_loader_for,
    load_document,
)

HAVE_REPORTLAB = importlib.util.find_spec("reportlab") is not None


# --------------------------------------------------------------------------- #
# Fixture builders (create real files at runtime with available deps)
# --------------------------------------------------------------------------- #


def _make_docx(path: Path) -> None:
    """Write a DOCX with a paragraph and a 3-row/2-col salary table."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("Employee Handbook")
    doc.add_paragraph("Compensation overview follows.")
    table = doc.add_table(rows=3, cols=2)
    data = [["Grade", "Salary"], ["A", "100000"], ["B", "90000"]]
    for r, row in enumerate(data):
        for c, value in enumerate(row):
            table.rows[r].cells[c].text = value
    doc.save(str(path))


def _make_xlsx(path: Path) -> None:
    """Write an XLSX whose single sheet holds a 3-row/2-col grid."""
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "Salaries"
    for row in [["Grade", "Salary"], ["A", 100000], ["B", 90000]]:
        ws.append(row)
    wb.save(str(path))


def _make_minimal_pdf(path: Path, lines: list[str]) -> None:
    """Build a tiny single-page PDF whose text is extractable by pypdf.

    Used only when reportlab is unavailable. The PDF contains a simple text
    stream (no ruled table geometry), which is sufficient to exercise PDF text
    extraction and loader dispatch.
    """

    def esc(s: str) -> str:
        return s.replace("\\", r"\\").replace("(", r"\(").replace(")", r"\)")

    content_lines = ["BT", "/F1 12 Tf", "72 720 Td", "14 TL"]
    for i, line in enumerate(lines):
        if i > 0:
            content_lines.append("T*")
        content_lines.append(f"({esc(line)}) Tj")
    content_lines.append("ET")
    content = ("\n".join(content_lines)).encode("latin-1")

    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Resources << /Font << /F1 5 0 R >> >> /Contents 4 0 R >>",
        b"<< /Length " + str(len(content)).encode() + b" >>\nstream\n" + content + b"\nendstream",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]

    out = b"%PDF-1.4\n"
    offsets = []
    for i, obj in enumerate(objects, start=1):
        offsets.append(len(out))
        out += f"{i} 0 obj\n".encode() + obj + b"\nendobj\n"
    xref_pos = len(out)
    n = len(objects) + 1
    out += f"xref\n0 {n}\n".encode()
    out += b"0000000000 65535 f \n"
    for off in offsets:
        out += f"{off:010d} 00000 n \n".encode()
    out += b"trailer\n"
    out += f"<< /Size {n} /Root 1 0 R >>\n".encode()
    out += b"startxref\n"
    out += f"{xref_pos}\n".encode()
    out += b"%%EOF"
    path.write_bytes(out)


def _make_pdf(path: Path) -> bool:
    """Create a PDF fixture. Returns True if it contains a ruled table.

    Prefers reportlab (real embedded table) when available; otherwise falls
    back to a hand-built text-only PDF.
    """
    if HAVE_REPORTLAB:
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Table
        from reportlab.lib.styles import getSampleStyleSheet

        styles = getSampleStyleSheet()
        doc = SimpleDocTemplate(str(path), pagesize=letter)
        data = [["Grade", "Salary"], ["A", "100000"], ["B", "90000"]]
        doc.build(
            [
                Paragraph("Employee Salary Report", styles["Title"]),
                Table(data),
            ]
        )
        return True

    _make_minimal_pdf(
        path, ["Employee Salary Report", "Grade A 100000", "Grade B 90000"]
    )
    return False


# --------------------------------------------------------------------------- #
# Loader dispatch / registry
# --------------------------------------------------------------------------- #


def test_supported_extensions_cover_all_formats():
    assert SUPPORTED_EXTENSIONS == {".pdf", ".docx", ".xlsx", ".txt", ".md"}


@pytest.mark.parametrize(
    "ext, cls",
    [
        (".pdf", PdfLoader),
        (".docx", DocxLoader),
        (".xlsx", XlsxLoader),
        (".txt", TextLoader),
        (".md", MarkdownLoader),
    ],
)
def test_get_loader_for_dispatches_by_extension(ext, cls):
    loader = get_loader_for(ext)
    assert isinstance(loader, cls)
    # All loaders satisfy the runtime-checkable Loader protocol.
    assert isinstance(loader, Loader)


def test_get_loader_for_normalizes_case_and_bare_extension_and_path():
    assert isinstance(get_loader_for(".PDF"), PdfLoader)
    assert isinstance(get_loader_for("pdf"), PdfLoader)
    assert isinstance(get_loader_for(Path("/tmp/report.DOCX")), DocxLoader)
    assert isinstance(get_loader_for("report.xlsx"), XlsxLoader)


def test_find_loader_returns_none_for_unsupported():
    assert find_loader(".exe") is None
    assert find_loader("archive.zip") is None


def test_get_loader_for_unsupported_raises():
    with pytest.raises(UnsupportedFormatError) as exc:
        get_loader_for(".csv")
    assert ".csv" in str(exc.value)
    # Message names supported formats so the processor can surface them.
    assert ".pdf" in str(exc.value)


def test_load_document_unsupported_extension_raises(tmp_path):
    bad = tmp_path / "data.csv"
    bad.write_text("a,b,c\n1,2,3\n", encoding="utf-8")
    with pytest.raises(UnsupportedFormatError):
        load_document(bad)


# --------------------------------------------------------------------------- #
# Nominal extraction paths
# --------------------------------------------------------------------------- #


def test_docx_loader_extracts_text_and_table(tmp_path):
    path = tmp_path / "handbook.docx"
    _make_docx(path)

    content = load_document(path)
    assert isinstance(content, ExtractedContent)
    assert "Employee Handbook" in content.text
    assert "Compensation overview follows." in content.text

    assert len(content.tables) == 1
    table = content.tables[0]
    # Row/column structure preserved: 3 rows x 2 cols.
    assert len(table) == 3
    assert all(len(row) == 2 for row in table)
    assert table[0] == ["Grade", "Salary"]
    assert table[1] == ["A", "100000"]
    assert table[2] == ["B", "90000"]


def test_xlsx_loader_extracts_sheet_as_table(tmp_path):
    path = tmp_path / "salaries.xlsx"
    _make_xlsx(path)

    content = XlsxLoader().load(path)
    assert content.text == ""
    assert len(content.tables) == 1
    table = content.tables[0]
    assert len(table) == 3
    assert all(len(row) == 2 for row in table)
    assert table[0] == ["Grade", "Salary"]
    # Numeric cells are stringified.
    assert table[1] == ["A", "100000"]
    assert table[2] == ["B", "90000"]


def test_txt_loader_preserves_plain_text(tmp_path):
    path = tmp_path / "notes.txt"
    body = "Line one\nLine two\n\nParagraph two."
    path.write_text(body, encoding="utf-8")

    content = load_document(path)
    assert content.text == body
    assert content.tables == []


def test_markdown_loader_preserves_source_as_is(tmp_path):
    path = tmp_path / "doc.md"
    body = "# Title\n\n| Grade | Salary |\n|---|---|\n| A | 100000 |\n"
    path.write_text(body, encoding="utf-8")

    content = load_document(path)
    # Markdown is preserved verbatim (tables stay inline in text).
    assert content.text == body
    assert content.tables == []


def test_pdf_loader_extracts_text_and_dispatches(tmp_path):
    path = tmp_path / "salary.pdf"
    has_ruled_table = _make_pdf(path)

    content = load_document(path)
    assert isinstance(content, ExtractedContent)
    assert "Salary Report" in content.text

    if has_ruled_table:
        # reportlab drew a real ruled table pdfplumber can recover.
        assert len(content.tables) >= 1
        flat = [cell for row in content.tables[0] for cell in row]
        assert "Grade" in flat and "Salary" in flat
    else:
        pytest.skip(
            "reportlab not installed: PDF table-extraction assertion skipped; "
            "PDF text extraction and loader dispatch still verified"
        )


# --------------------------------------------------------------------------- #
# Error paths
# --------------------------------------------------------------------------- #


def test_corrupt_docx_raises_parse_error(tmp_path):
    path = tmp_path / "broken.docx"
    path.write_bytes(b"this is not a valid docx / zip archive")
    with pytest.raises(DocumentParseError):
        load_document(path)


def test_corrupt_xlsx_raises_parse_error(tmp_path):
    path = tmp_path / "broken.xlsx"
    path.write_bytes(b"not a real xlsx workbook")
    with pytest.raises(DocumentParseError):
        load_document(path)


def test_corrupt_pdf_raises_parse_error(tmp_path):
    path = tmp_path / "broken.pdf"
    path.write_bytes(b"%PDF-1.4 garbage-not-a-real-pdf")
    with pytest.raises(DocumentParseError):
        load_document(path)
