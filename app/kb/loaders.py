"""Per-format document loaders (LangChain-style loader pattern).

Each supported upload format has one :class:`Loader` implementation that does
deterministic extraction of raw body text and tables from a file on disk,
returning an :class:`app.core.models.ExtractedContent`. Tables are returned as
a list of rows, each row a list of cell strings, so their row/column structure
survives before the AI-structuring step renders them as markdown (Req 5.2).

The loaders here intentionally do *no* AI work and *no* chunking — they are the
first, deterministic stage of the ingestion pipeline (see the design's
"Document Ingestion Pipeline"). The document processor selects a loader with
:func:`get_loader_for` (or the convenience :func:`load_document`) based on the
file extension; an unsupported extension raises :class:`UnsupportedFormatError`
so the processor can mark the document Status Error (Req 5.8), and a corrupt or
malformed file raises :class:`DocumentParseError` so the processor can mark
Status Error (Req 5.4).

Supported extensions: ``.pdf``, ``.docx``, ``.xlsx``, ``.txt``, ``.md``.

Design references:
* PDF: ``pypdf`` for text, ``pdfplumber`` for table cell grids.
* DOCX: ``python-docx`` paragraphs + tables.
* XLSX: ``openpyxl`` sheets rendered as tables.
* TXT / Markdown: plain text preserved as-is.
"""

from __future__ import annotations

from pathlib import Path
from typing import Protocol, runtime_checkable

from app.core.models import ExtractedContent

__all__ = [
    "Loader",
    "LoaderError",
    "UnsupportedFormatError",
    "DocumentParseError",
    "PdfLoader",
    "DocxLoader",
    "XlsxLoader",
    "TextLoader",
    "MarkdownLoader",
    "LOADERS",
    "SUPPORTED_EXTENSIONS",
    "get_loader_for",
    "find_loader",
    "load_document",
]


class LoaderError(Exception):
    """Base class for all loader errors raised by this module."""


class UnsupportedFormatError(LoaderError):
    """Raised when no loader is registered for a file extension.

    Carries the offending extension so callers can surface a clear message
    naming the supported formats (Req 5.8).
    """

    def __init__(self, extension: str) -> None:
        self.extension = extension
        supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        super().__init__(
            f"Unsupported document format {extension!r}; "
            f"supported formats are: {supported}"
        )


class DocumentParseError(LoaderError):
    """Raised when a file cannot be parsed by its loader.

    Wraps the underlying library exception so the document processor can mark
    the document Status Error with a clean message (Req 5.4) without leaking
    library-specific exception types.
    """

    def __init__(self, path: Path, reason: str) -> None:
        self.path = path
        self.reason = reason
        super().__init__(f"Failed to parse {path.name}: {reason}")


@runtime_checkable
class Loader(Protocol):
    """Common interface for a single-format document loader.

    Implementations declare the file extensions they handle and extract raw
    content from a file on disk. This mirrors the LangChain document-loader
    pattern: one loader per format behind a shared, mockable interface.

    Attributes:
        extensions: The lowercase file extensions (including the leading dot)
            this loader handles, e.g. ``{".pdf"}``.
    """

    extensions: set[str]

    def load(self, path: Path) -> ExtractedContent:
        """Extract text and tables from ``path``.

        Args:
            path: Filesystem path to the document to load.

        Returns:
            The extracted body text and tables.

        Raises:
            DocumentParseError: If the file cannot be parsed.
        """
        ...


def _normalize_cell(value: object) -> str:
    """Coerce a raw table cell value to a stripped string.

    Args:
        value: The raw cell value (may be ``None`` or a non-string type).

    Returns:
        The cell rendered as a string; ``None`` becomes an empty string.
    """
    if value is None:
        return ""
    return str(value).strip()


class PdfLoader:
    """Loader for PDF documents.

    Uses ``pypdf`` to extract body text page by page and ``pdfplumber`` to
    recover table cell grids so row/column structure is preserved (Req 5.2).
    """

    extensions: set[str] = {".pdf"}

    def load(self, path: Path) -> ExtractedContent:
        """Extract text (pypdf) and tables (pdfplumber) from a PDF.

        Args:
            path: Path to the ``.pdf`` file.

        Returns:
            Extracted text joined across pages and every detected table.

        Raises:
            DocumentParseError: If the PDF cannot be opened or parsed.
        """
        import pdfplumber
        from pypdf import PdfReader

        text_parts: list[str] = []
        tables: list[list[list[str]]] = []

        try:
            reader = PdfReader(str(path))
            for page in reader.pages:
                page_text = page.extract_text() or ""
                if page_text.strip():
                    text_parts.append(page_text.strip())
        except Exception as exc:  # noqa: BLE001 - normalize to DocumentParseError
            raise DocumentParseError(path, f"pypdf could not read PDF: {exc}") from exc

        try:
            with pdfplumber.open(str(path)) as pdf:
                for page in pdf.pages:
                    for raw_table in page.extract_tables() or []:
                        table = [
                            [_normalize_cell(cell) for cell in row]
                            for row in raw_table
                            if row is not None
                        ]
                        if table:
                            tables.append(table)
        except Exception as exc:  # noqa: BLE001 - normalize to DocumentParseError
            raise DocumentParseError(
                path, f"pdfplumber could not read PDF tables: {exc}"
            ) from exc

        return ExtractedContent(text="\n\n".join(text_parts), tables=tables)


class DocxLoader:
    """Loader for Word (.docx) documents using ``python-docx``.

    Extracts paragraph text as the body and each Word table as a preserved
    row/column grid (Req 5.2).
    """

    extensions: set[str] = {".docx"}

    def load(self, path: Path) -> ExtractedContent:
        """Extract paragraphs and tables from a DOCX file.

        Args:
            path: Path to the ``.docx`` file.

        Returns:
            Extracted paragraph text and every table in the document.

        Raises:
            DocumentParseError: If the document cannot be opened or parsed.
        """
        from docx import Document

        try:
            document = Document(str(path))
        except Exception as exc:  # noqa: BLE001 - normalize to DocumentParseError
            raise DocumentParseError(
                path, f"python-docx could not read document: {exc}"
            ) from exc

        text_parts = [
            para.text.strip()
            for para in document.paragraphs
            if para.text and para.text.strip()
        ]

        tables: list[list[list[str]]] = []
        for table in document.tables:
            grid = [
                [_normalize_cell(cell.text) for cell in row.cells]
                for row in table.rows
            ]
            if grid:
                tables.append(grid)

        return ExtractedContent(text="\n".join(text_parts), tables=tables)


class XlsxLoader:
    """Loader for Excel (.xlsx) workbooks using ``openpyxl``.

    Each worksheet is rendered as one table (a list of rows of cell strings),
    preserving the sheet's row/column structure (Req 5.2). Trailing fully empty
    rows/columns produced by openpyxl are trimmed.
    """

    extensions: set[str] = {".xlsx"}

    def load(self, path: Path) -> ExtractedContent:
        """Extract each worksheet of an XLSX workbook as a table.

        Args:
            path: Path to the ``.xlsx`` file.

        Returns:
            Extracted content whose ``tables`` holds one grid per non-empty
            worksheet; ``text`` is empty (spreadsheets carry no body prose).

        Raises:
            DocumentParseError: If the workbook cannot be opened or parsed.
        """
        from openpyxl import load_workbook

        try:
            workbook = load_workbook(filename=str(path), read_only=True, data_only=True)
        except Exception as exc:  # noqa: BLE001 - normalize to DocumentParseError
            raise DocumentParseError(
                path, f"openpyxl could not read workbook: {exc}"
            ) from exc

        tables: list[list[list[str]]] = []
        try:
            for sheet in workbook.worksheets:
                grid = [
                    [_normalize_cell(cell) for cell in row]
                    for row in sheet.iter_rows(values_only=True)
                ]
                grid = _trim_empty_edges(grid)
                if grid:
                    tables.append(grid)
        finally:
            workbook.close()

        return ExtractedContent(text="", tables=tables)


def _trim_empty_edges(grid: list[list[str]]) -> list[list[str]]:
    """Drop trailing all-empty rows and columns from a table grid.

    ``openpyxl`` read-only iteration can report empty padding rows/columns; this
    trims them so the returned grid reflects the sheet's real extent.

    Args:
        grid: A table as a list of rows of cell strings.

    Returns:
        The grid with trailing empty rows and columns removed. Returns an empty
        list if every cell is empty.
    """
    # Trim trailing empty rows.
    while grid and all(cell == "" for cell in grid[-1]):
        grid = grid[:-1]
    if not grid:
        return []
    # Trim trailing empty columns.
    width = max(len(row) for row in grid)
    normalized = [row + [""] * (width - len(row)) for row in grid]
    while width > 0 and all(row[width - 1] == "" for row in normalized):
        width -= 1
        normalized = [row[:width] for row in normalized]
    return [row for row in normalized] if width > 0 else []


class TextLoader:
    """Loader for plain-text (.txt) files.

    Reads the file as UTF-8 text (with a lenient fallback) and returns it as the
    body with no tables.
    """

    extensions: set[str] = {".txt"}

    def load(self, path: Path) -> ExtractedContent:
        """Read a plain-text file into body text.

        Args:
            path: Path to the ``.txt`` file.

        Returns:
            Extracted content with the full file text and no tables.

        Raises:
            DocumentParseError: If the file cannot be read.
        """
        return ExtractedContent(text=_read_text_file(path), tables=[])


class MarkdownLoader:
    """Loader for Markdown (.md) files.

    Preserves the markdown source as-is (Req 5.2): markdown tables and headings
    already carry their own structure, so no reformatting is done here.
    """

    extensions: set[str] = {".md"}

    def load(self, path: Path) -> ExtractedContent:
        """Read a Markdown file, preserving its source text verbatim.

        Args:
            path: Path to the ``.md`` file.

        Returns:
            Extracted content with the raw markdown as text and no tables.

        Raises:
            DocumentParseError: If the file cannot be read.
        """
        return ExtractedContent(text=_read_text_file(path), tables=[])


def _read_text_file(path: Path) -> str:
    """Read a text file as UTF-8, falling back to latin-1 on decode errors.

    Args:
        path: Path to the text file.

    Returns:
        The file contents as a string.

    Raises:
        DocumentParseError: If the file cannot be read from disk.
    """
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        try:
            return path.read_text(encoding="latin-1")
        except Exception as exc:  # noqa: BLE001 - normalize to DocumentParseError
            raise DocumentParseError(path, f"could not decode text file: {exc}") from exc
    except Exception as exc:  # noqa: BLE001 - normalize to DocumentParseError
        raise DocumentParseError(path, f"could not read text file: {exc}") from exc


# Registry of loader instances. Each extension maps to exactly one loader.
LOADERS: dict[str, Loader] = {}
for _loader in (PdfLoader(), DocxLoader(), XlsxLoader(), TextLoader(), MarkdownLoader()):
    for _ext in _loader.extensions:
        LOADERS[_ext] = _loader

#: All extensions covered by a registered loader (lowercase, dotted).
SUPPORTED_EXTENSIONS: set[str] = set(LOADERS)


def _extension_of(extension_or_path: str | Path) -> str:
    """Normalize an extension or path to a lowercase dotted extension.

    Args:
        extension_or_path: Either an extension (``"pdf"`` or ``".PDF"``) or a
            path/filename whose suffix is used.

    Returns:
        The lowercase extension including a leading dot, e.g. ``".pdf"``.
    """
    if isinstance(extension_or_path, Path):
        return extension_or_path.suffix.lower()
    value = str(extension_or_path)
    if "." in value and not value.startswith(".") and "/" not in value and "\\" not in value:
        # Bare "name.ext" style — take the suffix.
        value = Path(value).suffix
    ext = value.lower()
    if not ext.startswith("."):
        ext = "." + ext
    return ext


def find_loader(extension_or_path: str | Path) -> Loader | None:
    """Return the loader for an extension/path, or ``None`` if unsupported.

    Non-raising counterpart to :func:`get_loader_for`.

    Args:
        extension_or_path: An extension (``".pdf"``/``"pdf"``) or a path.

    Returns:
        The matching :class:`Loader`, or ``None`` when the format is not
        supported.
    """
    return LOADERS.get(_extension_of(extension_or_path))


def get_loader_for(extension_or_path: str | Path) -> Loader:
    """Return the loader for an extension/path.

    Args:
        extension_or_path: An extension (``".pdf"``/``"pdf"``) or a path.

    Returns:
        The matching :class:`Loader`.

    Raises:
        UnsupportedFormatError: If no loader handles the extension.
    """
    ext = _extension_of(extension_or_path)
    loader = LOADERS.get(ext)
    if loader is None:
        raise UnsupportedFormatError(ext)
    return loader


def load_document(path: str | Path) -> ExtractedContent:
    """Dispatch to the right loader by extension and extract content.

    Args:
        path: Path to the document to load.

    Returns:
        The extracted text and tables.

    Raises:
        UnsupportedFormatError: If the file's extension has no loader.
        DocumentParseError: If the file cannot be parsed.
    """
    p = Path(path)
    loader = get_loader_for(p)
    return loader.load(p)
