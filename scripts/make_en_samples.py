#!/usr/bin/env python3
"""Generate the English sample documents under ``samples/en``.
Produces one document per Intent_Space, each in a different supported format
(DOCX / XLSX / TXT / PDF) so every loader is exercised. All facts mirror the
Chinese corpus so answers stay consistent across languages. Idempotent.
"""

from __future__ import annotations

from pathlib import Path

EN_DIR = Path(__file__).resolve().parent.parent / "samples" / "en"


def make_docx() -> None:
    """HR employee handbook (DOCX) with the annual-leave table."""
    from docx import Document

    doc = Document()
    doc.add_heading("Zhilian Tech Co., Ltd. - Employee Handbook (2024)", 0)
    doc.add_paragraph(
        "Standard working hours are 9:00-18:00, Monday to Friday, with lunch "
        "break 12:00-13:00. Flexible arrival between 8:00 and 10:00 requires "
        "manager approval."
    )
    doc.add_heading("Annual leave by tenure", level=1)
    rows = [
        ("Tenure", "Annual leave days"), ("0-2 years", "10"),
        ("3-5 years", "12"), ("6-9 years", "15"), ("10+ years", "20"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    for r, (a, b) in enumerate(rows):
        table.rows[r].cells[0].text = a
        table.rows[r].cells[1].text = b
    doc.add_paragraph(
        "Up to 5 unused annual leave days may be carried over and must be used "
        "by March 31 of the following year. Employees receive 10 paid sick "
        "leave days per calendar year; a doctor's note is required for sick "
        "leave longer than 3 consecutive days. Paternity leave is 15 days."
    )
    doc.save(EN_DIR / "hr_employee_handbook.docx")


def make_xlsx() -> None:
    """Finance travel rates workbook (XLSX) with lodging and meal sheets."""
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "Lodging Rates"
    ws.append(["City Tier", "P1-P3 (CNY/night)", "P4-P5 (CNY/night)", "P6+ (CNY/night)"])
    for row in (
        ("Tier-1 (Beijing/Shanghai/Guangzhou/Shenzhen)", 450, 600, 800),
        ("Tier-2", 350, 450, 600),
        ("Other cities", 280, 350, 450),
        ("International (USD/night)", 120, 160, 220),
    ):
        ws.append(row)
    meals = wb.create_sheet("Meal Allowance")
    meals.append(["Trip Type", "Allowance"])
    meals.append(["Domestic travel", "150 CNY per day, no receipts required"])
    meals.append(["International travel", "60 USD per day"])
    wb.save(EN_DIR / "finance_travel_rates.xlsx")


def make_txt() -> None:
    """General IT support FAQ (plain text)."""
    text = """Zhilian Tech Co., Ltd. - IT Support FAQ (2024, English Edition)

Q: How do I contact the IT helpdesk?
A: Extension 8100 or it-helpdesk@zhilian.example.com, weekdays 9:00-18:00.
   Standard ticket response time is 4 hours.

Q: I forgot my domain account password. What should I do?
A: Reset it on the self-service password portal. At most 3 resets per day.

Q: How long can a VPN session last?
A: A single VPN session lasts at most 12 hours. Download the client from the
   internal software center and sign in with domain account + dynamic token.
   Never share your VPN account with anyone.

Q: How long can I book a meeting room for?
A: 2 hours maximum per booking. A room not used within 15 minutes of the
   start time is automatically released.

Q: I lost my employee badge.
A: Report the loss at the front desk (extension 8000). Replacement costs
   20 CNY and is issued within 1 business day.
"""
    (EN_DIR / "general_it_support_faq.txt").write_text(text, encoding="utf-8")


def make_pdf() -> None:
    """Legal contract approval policy as a minimal single-page PDF."""
    lines = [
        "Zhilian Tech Co., Ltd. - Contract Approval Policy (2024)",
        "",
        "Approval and signing authority by contract value:",
        "  Below 100,000 CNY: approved and signed by the department manager.",
        "  100,000 - 500,000 CNY: department director approval plus legal",
        "    review; signed by the department director.",
        "  500,000 - 2,000,000 CNY: Legal Director and Finance Director",
        "    approval; signed by the Deputy General Manager.",
        "  Above 2,000,000 CNY: General Manager office meeting approval;",
        "    signed by the legal representative.",
        "",
        "Every contract requires legal review before signing. Standard legal",
        "review takes 3 business days; an expedited 1-business-day channel",
        "requires department director approval.",
        "",
        "Disputes are submitted to the Beijing Arbitration Commission.",
        "On receiving a lawyer's letter, notify the Legal Department within",
        "24 hours; do not reply on your own.",
    ]
    content = "BT /F1 11 Tf 50 770 Td 15 TL\n"
    for line in lines:
        esc = line.replace("\\", r"\\").replace("(", r"\(").replace(")", r"\)")
        content += f"({esc}) Tj T*\n"
    stream = (content + "ET").encode("latin-1")
    objs = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        b"<< /Length %d >>\nstream\n" % len(stream) + stream + b"\nendstream",
    ]
    out = b"%PDF-1.4\n"
    offsets = []
    for i, obj in enumerate(objs, start=1):
        offsets.append(len(out))
        out += f"{i} 0 obj\n".encode() + obj + b"\nendobj\n"
    xref = len(out)
    out += b"xref\n0 6\n0000000000 65535 f \n"
    out += b"".join(f"{off:010d} 00000 n \n".encode() for off in offsets)
    out += b"trailer\n<< /Size 6 /Root 1 0 R >>\nstartxref\n%d\n%%%%EOF" % xref
    (EN_DIR / "legal_contract_approval_policy.pdf").write_bytes(out)


def main() -> None:
    """Generate all English sample documents."""
    EN_DIR.mkdir(parents=True, exist_ok=True)
    for make in (make_docx, make_xlsx, make_txt, make_pdf):
        make()
    for path in sorted(EN_DIR.iterdir()):
        print(f"wrote {path} ({path.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
